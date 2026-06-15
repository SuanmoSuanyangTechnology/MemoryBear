import logging
import re
from functools import reduce
from io import BytesIO

from docx import Document
from docx.image.exceptions import InvalidImageStreamError, UnexpectedEndOfFileError, UnrecognizedImageError
from docx.opc.oxml import parse_xml
from docx.opc.pkgreader import _SerializedRelationship, _SerializedRelationships
from PIL import Image

from app.core.rag.deepdoc.parser import DocxParser as RAGDocxParser
from app.core.rag.chunk.parser.base import DocumentParser
from app.core.rag.nlp import concat_img


def load_from_xml_v2(baseURI, rels_item_xml):
    srels = _SerializedRelationships()
    if rels_item_xml is not None:
        rels_elm = parse_xml(rels_item_xml)
        for rel_elm in rels_elm.Relationship_lst:
            if rel_elm.target_ref in ("../NULL", "NULL"):
                continue
            srels._srels.append(_SerializedRelationship(baseURI, rel_elm))
    return srels


class DocxParser(RAGDocxParser, DocumentParser):
    def parse(self, ctx):
        _SerializedRelationships.load_from_xml = load_from_xml_v2
        return self(ctx.filename, ctx.binary, ctx.from_page, ctx.to_page)

    def get_picture(self, document, paragraph):
        imgs = paragraph._element.xpath(".//pic:pic")
        if not imgs:
            return None
        res_img = None
        for img in imgs:
            embed = img.xpath(".//a:blip/@r:embed")
            if not embed:
                continue
            embed = embed[0]
            try:
                related_part = document.part.related_parts[embed]
                image_blob = related_part.image.blob
            except UnrecognizedImageError:
                logging.info("Unrecognized image format. Skipping image.")
                continue
            except UnexpectedEndOfFileError:
                logging.info("EOF was unexpectedly encountered while reading an image stream. Skipping image.")
                continue
            except InvalidImageStreamError:
                logging.info("The recognized image stream appears to be corrupted. Skipping image.")
                continue
            except UnicodeDecodeError:
                logging.info("The recognized image stream appears to be corrupted. Skipping image.")
                continue
            except Exception:
                logging.info("The recognized image stream appears to be corrupted. Skipping image.")
                continue
            try:
                image = Image.open(BytesIO(image_blob)).convert("RGB")
                if res_img is None:
                    res_img = image
                else:
                    res_img = concat_img(res_img, image)
            except Exception:
                continue

        return res_img

    def __clean(self, line):
        line = re.sub(r"\u3000", " ", line).strip()
        return line

    def __get_nearest_title(self, table_index, filename):
        import re
        from docx.text.paragraph import Paragraph

        titles = []
        blocks = []

        doc_name = re.sub(r"\.[a-zA-Z]+$", "", filename)
        if not doc_name:
            doc_name = "Untitled Document"

        try:
            for i, block in enumerate(self.doc._element.body):
                if block.tag.endswith("p"):
                    paragraph = Paragraph(block, self.doc)
                    blocks.append(("p", i, paragraph))
                elif block.tag.endswith("tbl"):
                    blocks.append(("t", i, None))
        except Exception as exc:
            logging.error(f"Error collecting blocks: {exc}")
            return ""

        target_table_pos = -1
        table_count = 0
        for _, (block_type, pos, _) in enumerate(blocks):
            if block_type == "t":
                if table_count == table_index:
                    target_table_pos = pos
                    break
                table_count += 1

        if target_table_pos == -1:
            return ""

        nearest_title = None
        for index in range(len(blocks) - 1, -1, -1):
            block_type, pos, block = blocks[index]
            if pos >= target_table_pos:
                continue

            if block_type != "p":
                continue

            if block.style and block.style.name and re.search(r"Heading\s*(\d+)", block.style.name, re.I):
                try:
                    level_match = re.search(r"(\d+)", block.style.name)
                    if level_match:
                        level = int(level_match.group(1))
                        if level <= 7:
                            title_text = block.text.strip()
                            if title_text:
                                nearest_title = (level, title_text)
                                break
                except Exception as exc:
                    logging.error(f"Error parsing heading level: {exc}")

        if nearest_title:
            titles.append(nearest_title)
            current_level = nearest_title[0]

            while current_level > 1:
                found = False
                for index in range(len(blocks) - 1, -1, -1):
                    block_type, pos, block = blocks[index]
                    if pos >= target_table_pos:
                        continue

                    if block_type != "p":
                        continue

                    if block.style and re.search(r"Heading\s*(\d+)", block.style.name, re.I):
                        try:
                            level_match = re.search(r"(\d+)", block.style.name)
                            if level_match:
                                level = int(level_match.group(1))
                                if level < current_level:
                                    title_text = block.text.strip()
                                    if title_text:
                                        titles.append((level, title_text))
                                        current_level = level
                                        found = True
                                        break
                        except Exception as exc:
                            logging.error(f"Error parsing parent heading: {exc}")

                if not found:
                    break

            titles.sort(key=lambda x: x[0])
            hierarchy = [doc_name] + [title[1] for title in titles]
            return " > ".join(hierarchy)

        return ""

    def __call__(self, filename, binary=None, from_page=0, to_page=100000):
        self.doc = Document(filename) if not binary else Document(BytesIO(binary))
        pn = 0
        lines = []
        last_image = None
        for paragraph in self.doc.paragraphs:
            if pn > to_page:
                break
            if from_page <= pn < to_page:
                if paragraph.text.strip():
                    if paragraph.style and paragraph.style.name == "Caption":
                        former_image = None
                        if lines and lines[-1][1] and lines[-1][2] != "Caption":
                            former_image = lines[-1][1].pop()
                        elif last_image:
                            former_image = last_image
                            last_image = None
                        lines.append((self.__clean(paragraph.text), [former_image], paragraph.style.name))
                    else:
                        current_image = self.get_picture(self.doc, paragraph)
                        image_list = [current_image]
                        if last_image:
                            image_list.insert(0, last_image)
                            last_image = None
                        lines.append((self.__clean(paragraph.text), image_list, paragraph.style.name if paragraph.style else ""))
                else:
                    if current_image := self.get_picture(self.doc, paragraph):
                        if lines:
                            lines[-1][1].append(current_image)
                        else:
                            last_image = current_image
            for run in paragraph.runs:
                if "lastRenderedPageBreak" in run._element.xml:
                    pn += 1
                    continue
                if "w:br" in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1
        new_line = [(line[0], reduce(concat_img, line[1]) if line[1] else None) for line in lines]

        tables = []
        for table_index, table in enumerate(self.doc.tables):
            title = self.__get_nearest_title(table_index, filename)
            html = "<table>"
            if title:
                html += f"<caption>Table Location: {title}</caption>"
            for row in table.rows:
                html += "<tr>"
                cell_index = 0
                try:
                    while cell_index < len(row.cells):
                        span = 1
                        cell = row.cells[cell_index]
                        for next_index in range(cell_index + 1, len(row.cells)):
                            if cell.text == row.cells[next_index].text:
                                span += 1
                                cell_index = next_index
                            else:
                                break
                        cell_index += 1
                        html += f"<td>{cell.text}</td>" if span == 1 else f"<td colspan='{span}'>{cell.text}</td>"
                except Exception as exc:
                    logging.warning(f"Error parsing table, ignore: {exc}")
                html += "</tr>"
            html += "</table>"
            tables.append(((None, html), ""))
        return new_line, tables

    def to_markdown(self, filename=None, binary=None, inline_images: bool = True):
        import base64
        import uuid

        import mammoth
        from markdownify import markdownify

        docx_file = BytesIO(binary) if binary else open(filename, "rb")

        def _convert_image_to_base64(image):
            try:
                with image.open() as image_file:
                    image_bytes = image_file.read()
                encoded = base64.b64encode(image_bytes).decode("utf-8")
                base64_url = f"data:{image.content_type};base64,{encoded}"
                alt_name = f"img_{uuid.uuid4().hex[:8]}"
                return {"src": base64_url, "alt": alt_name}
            except Exception as exc:
                logging.warning(f"Failed to convert image to base64: {exc}")
                return {"src": "", "alt": "image"}

        try:
            if inline_images:
                result = mammoth.convert_to_html(docx_file, convert_image=mammoth.images.img_element(_convert_image_to_base64))
            else:
                result = mammoth.convert_to_html(docx_file)

            html = result.value
            return markdownify(html)
        finally:
            if not binary:
                docx_file.close()
